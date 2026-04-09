import base64
import datetime
import json
import os
import re
import unicodedata
from collections import defaultdict
from copy import deepcopy
from io import BytesIO

from docx import Document
from docxtpl import DocxTemplate
from flask import Flask, jsonify, render_template, request, send_file

app = Flask(__name__)
UPLOAD_FOLDER = "uploads"
PLANTILLA_RULES_JSON = "data/plantilla_rules.json"
H2H_TEMPLATE_CANDIDATES_BY_TYPE = {
    "PGP_RENOVACION": [
        "templates/plantilla_h2h_pgp.docx",
        "templates/plantilla_h2h.docx",
    ],
    "OPENSSL_RENOVACION": [
        "templates/plantilla_h2h_openssl.docx",
        "templates/plantilla_h2h.docx",
    ],
    "CAMBIO_CREDENCIALES": [
        "templates/plantilla_h2h_cambio_credenciales.docx",
        "templates/plantilla_h2h.docx",
    ],
}
H2H_TEMPLATE_FALLBACK = [
    "templates/plantilla_internas.docx",
]
H2H_DOWNLOAD_TEMPLATES_BY_TYPE = {
    "PGP_RENOVACION": [
        ("MIS", "templates/plantilla_h2h_pgp_mis.docx"),
        ("FCD", "templates/plantilla_h2h_pgp_fcd.docx"),
    ],
    "OPENSSL_RENOVACION": [
        ("MIS", "templates/plantilla_h2h_openssl_mis.docx"),
        ("FCD", "templates/plantilla_h2h_openssl_fcd.docx"),
    ],
    "CAMBIO_CREDENCIALES": [
        ("CAMBIO_CREDENCIALES", "templates/plantilla_h2h.docx"),
    ],
}
INTERNAS_TEMPLATE_MATRIX = {
    "WIND_TO_MAINF": {
        "CDPremium": "templates/plantilla_windows_mainframe.docx",
        "Peer-to-Peer": "templates/plantilla_windows_mainframe.docx",
    },
    "MAINF_TO_WIND": {
        "CDPremium": "templates/plantilla_mainframe_distribuido.docx",
        "Peer-to-Peer": "templates/plantilla_internas.docx",
    },
    "WIND_MAINF_TO_MAINF_WIND": {
        "CDPremium": "templates/plantillas_win-mainf_main-win.docx",
        "Peer-to-Peer": "templates/plantilla_internas_bidireccional.docx",
    },
    "WIND_TO_WIND": {
        "CDPremium": "templates/plantilla_internas.docx",
        "Peer-to-Peer": "templates/plantilla_internas.docx",
    },
    "WIND_TO_LINUX": {
        "CDPremium": "templates/plantilla_internas.docx",
        "Peer-to-Peer": "templates/plantilla_internas.docx",
    },
    "LINUX_TO_LINUX": {
        "CDPremium": "templates/plantilla_internas.docx",
        "Peer-to-Peer": "templates/plantilla_internas.docx",
    },
    "LINUX_TO_WIND": {
        "CDPremium": "templates/plantilla_linux_windows.docx",
        "Peer-to-Peer": "templates/plantilla_internas.docx",
    },
}

# Creamos subcarpetas para cada categorÃ­a
CATEGORIAS = ["internas", "h2h", "noh2h"]
for cat in CATEGORIAS:
    os.makedirs(os.path.join(UPLOAD_FOLDER, cat), exist_ok=True)

def now():
    return datetime.datetime.now()


def es_categoria_valida(categoria):
    return categoria in CATEGORIAS


def normalizar_clave_plantilla(value):
    if not isinstance(value, str):
        return ""
    texto = unicodedata.normalize("NFKD", value).encode("ascii", "ignore").decode("ascii")
    return re.sub(r"\s+", " ", texto).strip().lower()


def construir_mapa_plantillas_normalizado():
    mapa = {}
    for tipo_documento, modalidades in INTERNAS_TEMPLATE_MATRIX.items():
        for modalidad, path in modalidades.items():
            key = (normalizar_clave_plantilla(tipo_documento), normalizar_clave_plantilla(modalidad))
            mapa[key] = path
    return mapa


INTERNAS_TEMPLATE_MAP = construir_mapa_plantillas_normalizado()


def obtener_plantilla_internas(tipo_documento, modalidad):
    key = (normalizar_clave_plantilla(tipo_documento), normalizar_clave_plantilla(modalidad))
    return INTERNAS_TEMPLATE_MAP.get(key)


def generar_rules_por_ambiente(lista_transferencias, regla_base, usuario_ft):
    nuevas_reglas = []
    for idx, transferencia in enumerate(lista_transferencias, start=1):
        regla = deepcopy(regla_base)

        servidor = transferencia.get("servidor", "DESCONOCIDO")
        ruta_origen = transferencia.get("rutaOrigen", "")
        archivo_origen = transferencia.get("archivoOrigen", "")
        ruta_destino = transferencia.get("rutaDestino", "")

        regla["comments"] = f"REGLA DEL {usuario_ft}"
        regla["priority"] = idx
        regla["name"] = f"MR_{usuario_ft}_PSSFTMWC_{servidor}_{idx:02d}"
        regla["procName"] = f"PL_{usuario_ft}_{servidor}_PSSFTMWC_01.cdp"
        ruta_origen_norm = str(ruta_origen or "").replace("\\", "/").rstrip("/")
        archivo_origen_norm = str(archivo_origen or "").replace("\\", "/").lstrip("/")
        ruta_archivo = f"{ruta_origen_norm}/{archivo_origen_norm}" if ruta_origen_norm else archivo_origen_norm

        regla["filePathKeyvalues"] = f"MATCH|{ruta_origen}"
        regla["fileNameKeyvalues"] = f"MATCH|{ruta_archivo}"
        regla["procArgs"] = f"&F1=%FA_FILE_FOUND. &A1=%FA_NOT_PATH. &D1={ruta_destino}"
        regla["lastModified"] = now().strftime("%a %b %d %H:%M:%S COT %Y")
        nuevas_reglas.append(regla)

    return {"rules": nuevas_reglas}


def codificar_json_en_base64(payload):
    contenido = json.dumps(payload, indent=4, ensure_ascii=False).encode("utf-8")
    return base64.b64encode(contenido).decode("ascii")


def normalizar_servidor(valor):
    return str(valor or "").strip().upper()


def normalizar_ambiente(valor):
    texto = str(valor or "").strip().lower()
    correcciones = {
        "ã³": "o",
        "Ã³": "o",
        "ó": "o",
        "ã­": "i",
        "Ã­": "i",
        "í": "i",
        "ã±": "n",
        "Ã±": "n",
        "ñ": "n",
    }
    for origen, destino in correcciones.items():
        texto = texto.replace(origen, destino)
    texto = unicodedata.normalize("NFKD", texto).encode("ascii", "ignore").decode("ascii")
    if "cert" in texto:
        return "certificacion"
    if "prod" in texto:
        return "produccion"
    return texto


def obtener_servidores_normalizados(lista_transferencias):
    servidores = set()
    for transferencia in lista_transferencias:
        servidor = normalizar_servidor(transferencia.get("servidor"))
        if servidor:
            servidores.add(servidor)
    return servidores


def validar_balance_y_servidores(lista_transferencias, etiqueta):
    total_cert = sum(1 for t in lista_transferencias if normalizar_ambiente(t.get("ambiente")) == "certificacion")
    total_prod = sum(1 for t in lista_transferencias if normalizar_ambiente(t.get("ambiente")) == "produccion")
    if total_cert == 0 or total_prod == 0:
        return (
            "Debe existir al menos una transferencia de CertificaciÃ³n y una de ProducciÃ³n "
            f"en {etiqueta}."
        )
    if total_cert != total_prod:
        return (
            "Debe haber el mismo nÃºmero de transferencias en CertificaciÃ³n y ProducciÃ³n "
            f"en {etiqueta}. (CertificaciÃ³n: {total_cert} | ProducciÃ³n: {total_prod})"
        )

    servidores_cert = obtener_servidores_normalizados(
        [t for t in lista_transferencias if normalizar_ambiente(t.get("ambiente")) == "certificacion"]
    )
    servidores_prod = obtener_servidores_normalizados(
        [t for t in lista_transferencias if normalizar_ambiente(t.get("ambiente")) == "produccion"]
    )
    repetidos = sorted(servidores_cert.intersection(servidores_prod))
    if repetidos:
        return (
            "No se permite el mismo servidor en CertificaciÃ³n y ProducciÃ³n "
            f"en {etiqueta}. Repetidos: {', '.join(repetidos)}."
        )
    return ""


def combinar_transferencias_bidireccionales(transferencias_w2m, transferencias_m2w):
    combinadas = []
    for t in transferencias_w2m:
        combinadas.append(
            {
                "ambiente": t.get("ambiente", ""),
                "servidor": t.get("servidor", ""),
                "rutaOrigen": t.get("rutaOrigen", ""),
                "archivoOrigen": t.get("archivoOrigen", ""),
                "rutaDestino": t.get("dataset", ""),
                "archivoDestino": t.get("job", ""),
                "dataset": t.get("dataset", ""),
                "job": t.get("job", ""),
            }
        )
    for t in transferencias_m2w:
        nombre_archivo = t.get("nombreArchivo", "")
        combinadas.append(
            {
                "ambiente": t.get("ambiente", ""),
                "servidor": t.get("servidor", ""),
                "rutaOrigen": t.get("rutaOrigen", ""),
                "archivoOrigen": nombre_archivo,
                "rutaDestino": t.get("rutaDestino", ""),
                "archivoDestino": nombre_archivo,
                "dataset": t.get("rutaDestino", ""),
                "job": nombre_archivo,
            }
        )
    return combinadas


@app.route("/")
def index():
    return render_template("index.html")

@app.route("/<categoria>")
def listar_categoria(categoria):
    if not es_categoria_valida(categoria):
        return "Categoría no encontrada", 404

    plantillas = {
        "internas": "interfaz_internas.html",
        "h2h": "interfaz_h2h.html",
        "noh2h": "interfaz_noh2h.html",
    }
    return render_template(plantillas[categoria], tipo=categoria.upper())
def obtener_plantilla_h2h(tipo_documento):
    candidatos = H2H_TEMPLATE_CANDIDATES_BY_TYPE.get(tipo_documento, []) + H2H_TEMPLATE_FALLBACK
    for path in candidatos:
        if os.path.exists(path):
            return path
    return None


def obtener_plantillas_h2h_descarga(tipo_documento):
    plantillas = []
    for etiqueta, path in H2H_DOWNLOAD_TEMPLATES_BY_TYPE.get(tipo_documento, []):
        if os.path.exists(path):
            plantillas.append((etiqueta, path))
    if plantillas:
        return plantillas

    plantilla = obtener_plantilla_h2h(tipo_documento)
    if plantilla:
        return [("DOCUMENTO", plantilla)]
    return []


def construir_datos_h2h_desde_payload(data):
    usuarios = data.get("usuarios", [])
    usuarios_normalizados = []
    if isinstance(usuarios, list):
        for usuario in usuarios:
            if not isinstance(usuario, dict):
                continue
            usuarios_normalizados.append(
                {
                    "nombre_empresa": str(usuario.get("nombreEmpresa", "")).strip(),
                    "nombre_usuario": str(usuario.get("nombreUsuario", "")).strip(),
                    "nombre_certificado": str(usuario.get("nombreCertificado", "")).strip(),
                    "fecha_vencimiento": str(usuario.get("fechaVencimiento", "")).strip(),
                    "fecha_vencimiento_anterior": str(usuario.get("fechaVencimientoAnterior", "")).strip(),
                }
            )

    return {
        "tipo_documento": str(data.get("tipoDocumento", "")).strip().upper(),
        "nombre_mvp": str(data.get("nombreMVP", "")).strip(),
        "nombre_empresa": str(data.get("nombreEmpresa", "")).strip(),
        "nombre_usuario": str(data.get("nombreUsuario", "")).strip(),
        "nombre_certificado": str(data.get("nombreCertificado", "")).strip(),
        "fecha_vencimiento": str(data.get("fechaVencimiento", "")).strip(),
        "fecha_vencimiento_anterior": str(data.get("fechaVencimientoAnterior", "")).strip(),
        "commit": str(data.get("commit", "")).strip(),
        "ruta_bitbucket": str(data.get("rutaBitbucket", "")).strip(),
        "usuarios": usuarios_normalizados,
    }


def validar_datos_h2h(datos):
    tipo = datos.get("tipo_documento", "")
    if tipo not in {"PGP_RENOVACION", "OPENSSL_RENOVACION", "CAMBIO_CREDENCIALES"}:
        return "Tipo de documento no vÃ¡lido."

    requeridos_base = ["nombre_mvp"]
    for campo in requeridos_base:
        if not datos.get(campo):
            return f"El campo '{campo}' es obligatorio."

    if tipo in {"PGP_RENOVACION", "OPENSSL_RENOVACION"}:
        usuarios = datos.get("usuarios", [])
        if not usuarios:
            return "Debe agregar al menos un usuario para la renovacion."

        for indice, usuario in enumerate(usuarios, start=1):
            requeridos_usuario = ["nombre_empresa", "nombre_usuario", "nombre_certificado"]
            for campo in requeridos_usuario:
                if not usuario.get(campo):
                    return f"El campo '{campo}' es obligatorio en la fila {indice}."
            for campo_fecha in ["fecha_vencimiento", "fecha_vencimiento_anterior"]:
                if not usuario.get(campo_fecha):
                    return f"El campo '{campo_fecha}' es obligatorio en la fila {indice}."

    if tipo == "CAMBIO_CREDENCIALES":
        if not datos.get("nombre_empresa"):
            return "El campo 'nombre_empresa' es obligatorio para CAMBIO_CREDENCIALES."
        if not datos.get("nombre_usuario"):
            return "El campo 'nombre_usuario' es obligatorio para CAMBIO_CREDENCIALES."

    return ""


def generar_documento_h2h_fallback(datos):
    doc = Document()
    doc.add_heading("Generar documentos H2H", level=1)
    doc.add_paragraph(f"Fecha: {now().strftime('%d/%m/%Y %H:%M:%S')}")

    filas = [
        ("Tipo documento", datos.get("tipo_documento", "")),
        ("Nombre OCD", datos.get("nombre_mvp", "")),
        ("Nombre empresa", datos.get("nombre_empresa", "")),
        ("Nombre usuario FT", datos.get("nombre_usuario", "")),
        ("Nombre certificado", datos.get("nombre_certificado", "")),
        ("Fecha vencimiento", datos.get("fecha_vencimiento", "")),
        ("Fecha vencimiento anterior", datos.get("fecha_vencimiento_anterior", "")),
        ("Commit", datos.get("commit", "")),
        ("Ruta bitbucket", datos.get("ruta_bitbucket", "")),
    ]

    tabla = doc.add_table(rows=1, cols=2)
    tabla.style = "Table Grid"
    hdr = tabla.rows[0].cells
    hdr[0].text = "Campo"
    hdr[1].text = "Valor"
    for campo, valor in filas:
        row = tabla.add_row().cells
        row[0].text = campo
        row[1].text = str(valor or "")

    usuarios = datos.get("usuarios", [])
    if usuarios:
        doc.add_paragraph("")
        doc.add_paragraph("Usuarios")
        tabla_usuarios = doc.add_table(rows=1, cols=5)
        tabla_usuarios.style = "Table Grid"
        encabezados = tabla_usuarios.rows[0].cells
        encabezados[0].text = "Empresa"
        encabezados[1].text = "Usuario FT"
        encabezados[2].text = "Certificado"
        encabezados[3].text = "Fecha vencimiento"
        encabezados[4].text = "Fecha anterior"

        for usuario in usuarios:
            fila = tabla_usuarios.add_row().cells
            fila[0].text = str(usuario.get("nombre_empresa", "") or "")
            fila[1].text = str(usuario.get("nombre_usuario", "") or "")
            fila[2].text = str(usuario.get("nombre_certificado", "") or "")
            fila[3].text = str(usuario.get("fecha_vencimiento", "") or "")
            fila[4].text = str(usuario.get("fecha_vencimiento_anterior", "") or "")

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


def unir_usuarios_ft(usuarios):
    nombres = []
    for usuario in usuarios:
        nombre = str(usuario.get("nombre_usuario", "")).strip()
        if nombre:
            nombres.append(nombre)

    if not nombres:
        return ""
    if len(nombres) == 1:
        return nombres[0]
    if len(nombres) == 2:
        return f"{nombres[0]} y {nombres[1]}"
    return f"{', '.join(nombres[:-1])} y {nombres[-1]}"


def unir_empresas_usuarios(usuarios):
    pares = []
    for usuario in usuarios:
        empresa = str(usuario.get("nombre_empresa", "")).strip()
        nombre = str(usuario.get("nombre_usuario", "")).strip()
        if empresa and nombre:
            pares.append(f"{empresa} - {nombre}")

    return " | ".join(pares)


def renderizar_documento_h2h(plantilla_path, context, datos):
    try:
        doc = DocxTemplate(plantilla_path)
        doc.render(context)
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer
    except Exception as render_error:
        print(f"Plantilla H2H invalida ({plantilla_path}): {render_error}. Usando fallback.")
        return generar_documento_h2h_fallback(datos)


@app.route("/guardar_h2h", methods=["POST"])
def guardar_h2h():
    try:
        payload = request.get_json() or {}
        datos = construir_datos_h2h_desde_payload(payload)
        mensaje_error = validar_datos_h2h(datos)
        if mensaje_error:
            return jsonify({"success": False, "message": mensaje_error}), 400

        plantillas_descarga = obtener_plantillas_h2h_descarga(datos.get("tipo_documento", ""))
        if not plantillas_descarga:
            return (
                jsonify(
                    {
                        "success": False,
                        "message": (
                            "No se encontrÃ³ plantilla para H2H. "
                            "Agrega las plantillas en templates/."
                        ),
                    }
                ),
                404,
            )

        momento = now()
        usuarios = datos.get("usuarios", [])
        primer_usuario = usuarios[0] if usuarios else {}
        context = {
            **datos,
            "usuarios": usuarios,
            "total_usuarios": len(usuarios),
            "cantidad_usuario": len(usuarios),
            "usuarios_ft_texto": unir_usuarios_ft(usuarios),
            "usuarios_ft_texto_sin_primero": unir_usuarios_ft(usuarios[1:]),
            "empresas_usuarios_texto": unir_empresas_usuarios(usuarios),
            "empresas_usuarios_texto_sin_primero": unir_empresas_usuarios(usuarios[1:]),
            "nombre_empresa": datos.get("nombre_empresa") or primer_usuario.get("nombre_empresa", ""),
            "nombre_usuario": datos.get("nombre_usuario") or primer_usuario.get("nombre_usuario", ""),
            "nombre_certificado": datos.get("nombre_certificado") or primer_usuario.get("nombre_certificado", ""),
            "fecha_vencimiento": datos.get("fecha_vencimiento") or primer_usuario.get("fecha_vencimiento", ""),
            "fecha_vencimiento_anterior": (
                datos.get("fecha_vencimiento_anterior")
                or primer_usuario.get("fecha_vencimiento_anterior", "")
            ),
            "fecha_actual": momento.strftime("%d/%m/%Y"),
            "hora_actual": momento.strftime("%H:%M:%S"),
            "es_renovacion": datos["tipo_documento"] in {"OPENSSL_RENOVACION", "PGP_RENOVACION"},
            "es_pgp": datos["tipo_documento"] == "PGP_RENOVACION",
            "es_openssl": datos["tipo_documento"] == "OPENSSL_RENOVACION",
            "es_cambio_credenciales": datos["tipo_documento"] == "CAMBIO_CREDENCIALES",
        }

        nombre_mvp = re.sub(r'[\\/:*?"<>|]+', "_", datos.get("nombre_mvp") or "SIN_MVP")
        archivos = []
        for etiqueta, plantilla_path in plantillas_descarga:
            buffer = renderizar_documento_h2h(plantilla_path, context, datos)
            if etiqueta in {"MIS", "FCD"}:
                nombre_archivo = f"SSFT-{etiqueta}-{nombre_mvp}.docx"
            else:
                nombre_archivo = f"SSFT-H2H-{nombre_mvp}-{etiqueta}.docx"
            archivos.append(
                {
                    "filename": nombre_archivo,
                    "content_base64": base64.b64encode(buffer.getvalue()).decode("ascii"),
                    "mimetype": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                }
            )

        return jsonify(
            {
                "success": True,
                "files": archivos,
            }
        )
    except Exception as e:
        print(f"Error en guardar_h2h: {e}")
        return jsonify({"success": False, "message": f"OcurriÃ³ un error: {e}"}), 500


@app.route("/guardar_internas", methods=["POST"])
def guardar_internas():
    try:
        data = request.get_json() or {}

        tipo_documento = data.get("tipoDocumento", "")
        nombre_ocd = data.get("nombreOCD", "")
        usuario_ft = data.get("usuarioFT", "")
        modalidad = data.get("modalidad", "")
        transferencias_w2m = data.get("transferenciasW2M", [])
        transferencias_m2w = data.get("transferenciasM2W", [])
        transferencias = data.get("transferencias", [])
        if (
            tipo_documento == "WIND_MAINF_TO_MAINF_WIND"
            and (transferencias_w2m or transferencias_m2w)
            and not transferencias
        ):
            transferencias = combinar_transferencias_bidireccionales(transferencias_w2m, transferencias_m2w)
        transferencias = [
            {
                **t,
                "dataset": t.get("dataset", t.get("rutaDestino", "")),
                "job": t.get("job", t.get("archivoDestino", "")),
            }
            for t in transferencias
        ]
        if tipo_documento == "WIND_MAINF_TO_MAINF_WIND":
            if not transferencias_w2m or not transferencias_m2w:
                return (
                    jsonify(
                        {
                            "success": False,
                            "message": "Se requieren ambas secciones: WINDOW A MAINFRAME y MAINFRAME A WINDOWS.",
                        }
                    ),
                    400,
                )

            error_w2m = validar_balance_y_servidores(transferencias_w2m, "WINDOW A MAINFRAME")
            if error_w2m:
                return jsonify({"success": False, "message": error_w2m}), 400

            error_m2w = validar_balance_y_servidores(transferencias_m2w, "MAINFRAME A WINDOWS")
            if error_m2w:
                return jsonify({"success": False, "message": error_m2w}), 400

        transferencias_cert = [
            t for t in transferencias if normalizar_ambiente(t.get("ambiente")) == "certificacion"
        ]
        transferencias_prod = [
            t for t in transferencias if normalizar_ambiente(t.get("ambiente")) == "produccion"
        ]
        servidores_cert = obtener_servidores_normalizados(transferencias_cert)
        servidores_prod = obtener_servidores_normalizados(transferencias_prod)
        servidores_repetidos = sorted(servidores_cert.intersection(servidores_prod))

        if servidores_repetidos:
            return (
                jsonify(
                    {
                        "success": False,
                        "message": (
                            "No se permite el mismo servidor en CertificaciÃ³n y ProducciÃ³n. "
                            f"Repetidos: {', '.join(servidores_repetidos)}."
                        ),
                    }
                ),
                400,
            )

        with open(PLANTILLA_RULES_JSON, "r", encoding="utf-8") as f:
            rules_template = json.load(f)
        if "rules" not in rules_template or not rules_template["rules"]:
            raise Exception("La plantilla JSON no contiene reglas base")
        regla_base = rules_template["rules"][0]

        rules_cert = generar_rules_por_ambiente(transferencias_cert, regla_base, usuario_ft)
        rules_prod = generar_rules_por_ambiente(transferencias_prod, regla_base, usuario_ft)

        json_cert_b64 = codificar_json_en_base64(rules_cert)
        json_prod_b64 = codificar_json_en_base64(rules_prod)

        json_cert_name = f"SSFT-MIS-{nombre_ocd}-CERTIFICACION.json"
        json_prod_name = f"SSFT-MIS-{nombre_ocd}-PRODUCCION.json"

        plantilla_path = obtener_plantilla_internas(tipo_documento, modalidad)
        if not plantilla_path:
            return (
                jsonify(
                    {
                        "success": False,
                        "message": (
                            f"No hay plantilla configurada para tipo_documento='{tipo_documento}' "
                            f"y modalidad='{modalidad}'."
                        ),
                    }
                ),
                400,
            )

        if not os.path.exists(plantilla_path):
            return (
                jsonify(
                    {
                        "success": False,
                        "message": (
                            f"No existe la plantilla configurada para tipo_documento='{tipo_documento}' "
                            f"y modalidad='{modalidad}': {plantilla_path}"
                        ),
                    }
                ),
                404,
            )

        doc = DocxTemplate(plantilla_path)
        rutas_origen = sorted({t["rutaOrigen"] for t in transferencias if t.get("rutaOrigen")})

        certificacion_agrupada = agrupar_por_servidor(transferencias_cert)
        produccion_agrupada = agrupar_por_servidor(transferencias_prod)
        filas_watch_directory = construir_filas_watch_directory(certificacion_agrupada, produccion_agrupada)

        context = {
            "tipo_documento": tipo_documento,
            "nombre_ocd": nombre_ocd,
            "usuario_ft": usuario_ft,
            "modalidad": modalidad,
            "transferencias": transferencias,
            "transferencias_w2m": transferencias_w2m,
            "transferencias_m2w": transferencias_m2w,
            "certificacion": certificacion_agrupada,
            "produccion": produccion_agrupada,
            "rutas_origen": rutas_origen,
            "resumen_rutas_bidireccionales": construir_resumen_rutas_bidireccionales(
                transferencias_w2m, transferencias_m2w
            ),
            "filas_watch_directory": filas_watch_directory,
            "watch_directory_text": construir_texto_watch_directory(filas_watch_directory),
            "transferencias_text": construir_texto_transferencias(transferencias),
        }
        doc.render(context)

        output_filename = f"SSFT-MIS-{nombre_ocd}.docx"
        buf = BytesIO()
        doc.save(buf)
        buf.seek(0)

        response = send_file(
            buf,
            as_attachment=True,
            download_name=output_filename,
            mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
        # JSON VIAJA EN MEMORIA (NO DISCO)
        response.headers["X-JSON-CERT-NAME"] = json_cert_name
        response.headers["X-JSON-CERT-DATA"] = json_cert_b64

        response.headers["X-JSON-PROD-NAME"] = json_prod_name
        response.headers["X-JSON-PROD-DATA"] = json_prod_b64
        return response
    except Exception as e:
        print(f"Error en guardar_internas: {e}")
        return jsonify({"success": False, "message": f"OcurriÃ³ un error: {e}"}), 500



def agrupar_por_servidor(lista):
    grupos = defaultdict(lambda: defaultdict(list))

    for t in lista:
        servidor = t.get("servidor")
        ruta = t.get("rutaOrigen", "")
        if not servidor:
            continue

        ruta_norm = ruta.strip()

        grupos[servidor][ruta_norm].append(
            {
                "ambiente": t.get("ambiente"),
                "archivoOrigen": t.get("archivoOrigen"),
                "rutaOrigen": t.get("rutaOrigen"),
                "servidorDestino": t.get("servidorDestino", ""),
                "archivoDestino": t.get("archivoDestino"),
                "rutaDestino": t.get("rutaDestino"),
                "dataset": t.get("dataset", t.get("rutaDestino", "")),
                "job": t.get("job", t.get("archivoDestino", "")),
            }
        )

    servidores = []
    for servidor, rutas_map in grupos.items():
        rutas_list = []
        for ruta, transf_list in rutas_map.items():
            rutas_list.append({"rutaOrigen": ruta, "transferencias": transf_list})
        servidores.append({"servidor": servidor, "rutas": rutas_list})

    return servidores


def construir_resumen_rutas_bidireccionales(transferencias_w2m, transferencias_m2w):
    filas_unicas = {}

    for t in transferencias_w2m:
        ambiente = (t.get("ambiente") or "").strip()
        servidor = (t.get("servidor") or "").strip()
        ruta = (t.get("rutaOrigen") or "").strip()
        if not ambiente or not servidor or not ruta:
            continue
        key = (ambiente.upper(), servidor.upper(), ruta.upper())
        if key not in filas_unicas:
            filas_unicas[key] = {"ambiente": ambiente, "servidor": servidor, "ruta": ruta}

    for t in transferencias_m2w:
        ambiente = (t.get("ambiente") or "").strip()
        servidor = (t.get("servidor") or "").strip()
        ruta = (t.get("rutaDestino") or "").strip()
        if not ambiente or not servidor or not ruta:
            continue
        key = (ambiente.upper(), servidor.upper(), ruta.upper())
        if key not in filas_unicas:
            filas_unicas[key] = {"ambiente": ambiente, "servidor": servidor, "ruta": ruta}

    return sorted(
        filas_unicas.values(),
        key=lambda x: (
            x["servidor"].upper(),
            x["ruta"].upper(),
            x["ambiente"].upper(),
        ),
    )


def construir_filas_watch_directory(certificacion, produccion):
    filas = []
    vistos = set()
    for servidor_data in (certificacion or []) + (produccion or []):
        servidor = (servidor_data.get("servidor") or "").strip()
        for ruta_data in servidor_data.get("rutas", []):
            ruta_origen = (ruta_data.get("rutaOrigen") or "").strip()
            if not servidor or not ruta_origen:
                continue

            key = (servidor.upper(), ruta_origen.upper())
            if key in vistos:
                continue
            vistos.add(key)

            transferencias = ruta_data.get("transferencias") or []
            ambiente_raw = transferencias[0].get("ambiente", "") if transferencias else ""
            ambiente_norm = normalizar_ambiente(ambiente_raw)
            if ambiente_norm == "certificacion":
                ambiente = "Certificación"
            elif ambiente_norm == "produccion":
                ambiente = "Producción"
            else:
                ambiente = ambiente_raw

            filas.append(
                {
                    "ambiente": ambiente,
                    "node": f"CD{servidor}",
                    "watch_directory": ruta_origen,
                    "directory_description": "DIRECTORIO HACIA CDP",
                }
            )

    return sorted(
        filas,
        key=lambda x: (
            (x.get("node") or "").upper(),
            (x.get("watch_directory") or "").upper(),
            (x.get("ambiente") or "").upper(),
        ),
    )


def construir_texto_watch_directory(filas_watch_directory):
    if not filas_watch_directory:
        return "Sin registros."
    lineas = []
    for fila in filas_watch_directory:
        lineas.append(
            " | ".join(
                [
                    str(fila.get("ambiente", "")),
                    str(fila.get("node", "")),
                    str(fila.get("watch_directory", "")),
                    str(fila.get("directory_description", "")),
                ]
            )
        )
    return "\n".join(lineas)


def construir_texto_transferencias(transferencias):
    if not transferencias:
        return "Sin transferencias."
    lineas = []
    for t in transferencias:
        lineas.append(
            " | ".join(
                [
                    str(t.get("ambiente", "")),
                    str(t.get("servidor", "")),
                    str(t.get("rutaOrigen", "")),
                    str(t.get("archivoOrigen", "")),
                    str(t.get("rutaDestino", "")),
                    str(t.get("archivoDestino", "")),
                    str(t.get("servidorDestino", "")),
                ]
            )
        )
    return "\n".join(lineas)

if __name__ == "__main__":
    app.run(debug=True)

    
